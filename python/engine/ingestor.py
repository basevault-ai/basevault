"""
Ingestion layer — reads heterogeneous input formats and produces normalized Document objects.

Supported formats (whitelist — anything else is skipped):
  .txt  — plain text or WhatsApp export
  .md, .markdown
  .pdf  — text extraction only
  .docx
  .json — Day One export
  .zip  — recursively ingested (including nested zips)
  Images: .jpg .jpeg .png .heic .heif .bmp .tiff .webp .gif
          (transcribed via vision LLM, not OCR)

Explicit skips:
  .html / .htm — excluded to exercise the skip path
  Any file > 20 MB — dropped

Usage:
    from engine.ingestor import ingest
    docs = ingest("/path/to/file_or_dir")
"""
from __future__ import annotations

import json
import re
import shutil
import tempfile
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field

# `pypdf` + `python-docx` are imported lazily inside `_parse_pdf` /
# `_parse_docx` (their only call sites). Module-top imports dragged
# ~45 ms of dead weight onto every interactive chatbot turn, which
# pulls `retrieval → embeddings → rag_enricher → rag_chunker →
# ingestor` but never parses a PDF/DOCX. The ingestion path is
# behaviorally unchanged — same imports, just deferred to the two
# parsers that use them.


# Runner injects _log_write here via runner._patch_llm_calls so ingest
# warnings (skipped files, failed PDFs, etc.) land in run.log instead
# of being lost to stdout. Falls back to plain print() for CLI / test
# invocations that never touched the runner.
_runner_log = None


def _log(msg: str) -> None:
    if _runner_log is not None:
        try:
            _runner_log(msg)
            return
        except Exception:
            pass
    try:
        print(msg, flush=True)
    except BrokenPipeError:
        pass
from enum import Enum
from pathlib import Path

from engine.llm import Mode, max_workers


MAX_FILE_SIZE = 40 * 1024 * 1024  # 40 MB


class SourceType(str, Enum):
    WHATSAPP = "whatsapp"
    MD_FILE = "md_file"
    TXT = "txt"
    DAYONE_JSON = "dayone_json"
    PDF = "pdf"
    DOCX = "docx"
    IMAGE = "image"
    CLAUDE_WEB_CONVERSATION = "claude_web_conversation"
    CLAUDE_WEB_PROJECT = "claude_web_project"
    CLAUDE_CODE_SESSION = "claude_code_session"
    CLAUDE_CODE_HISTORY = "claude_code_history"
    CODEX_SESSION = "codex_session"
    CODEX_HISTORY = "codex_history"
    CHATGPT_CONVERSATION = "chatgpt_conversation"
    UNKNOWN = "unknown"


_IMAGE_EXTS = {
    ".jpg", ".jpeg", ".png", ".heic", ".heif",
    ".bmp", ".tiff", ".webp", ".gif",
}


# Rough size estimate (in bytes) of an image's transcribed text OR a
# PDF/DOCX's extracted text. Used in the pre-ingest stats-from-paths
# path so binary-format files (where on-disk size doesn't reflect
# text content) don't blow up the chunk-count estimate. A 50MB PDF
# might have ~10kB of actual extracted text; using its file size as
# a token proxy would 10–50× the estimate.
_PRE_SCAN_BINARY_TEXT_BYTES = 10_000


PipelineStats = dict  # type alias — see schema in docstrings below
"""Stats shape (the contract — pre and post ingestion):

  {
    "n_images":     int,        # vision est_calls
    "n_text_files": int,        # non-image documents
    "text_bytes":   int,        # sum of text content across ALL
                                # documents (image transcripts +
                                # text-file content + extracted
                                # PDF/DOCX text). Tokens proxy =
                                # text_bytes / 3.
    "n_splits":     int | None, # splitter output count (post-ingest
                                # only — None pre-ingest, since the
                                # splitter hasn't run yet). When
                                # present, drives the extract
                                # est_calls directly.
    "chunk_cap":    int,        # token cap per extract chunk
    "n_topics":     int,        # active topic taxonomy size
  }

Two producers, one shape:
  pipeline_stats_from_paths(paths, ...)  → pre-ingest, fake stats
                                            from file stat() walks
  pipeline_stats_from_docs(docs, n_splits, ...) → post-ingest,
                                            real stats from parsed
                                            Documents + splitter
                                            output

The estimator that consumes these dicts lives in `progress.py`
(`progress.estimate_pipeline`) — workload projection sits next to
`FALLBACK_SECONDS_PER_CALL` / `PARALLELISM_PER_STAGE` rather than
adjacent to its inputs. Adding a new producer (resume from
checkpoint, say) = write a function that returns the same dict shape.
"""


def pipeline_stats_from_paths(
    paths, chunk_cap: int = 30_000, n_topics: int | None = None,
) -> dict:
    """Pre-ingest pipeline stats. Walks `paths` (stat + zip
    namelist; no file content read) and returns a `PipelineStats`
    dict. Synthesis rules:
      - Image / PDF / DOCX: contributes `_PRE_SCAN_BINARY_TEXT_BYTES`
        (10kB) to `text_bytes`. File size on disk doesn't reflect
        text content for these formats.
      - Text formats (md / txt / json): real on-disk byte size.
      - Zips: peek namelist + member sizes recursively.
    `n_splits` is None — splitter hasn't run yet.
    """
    from pathlib import Path
    n_images = 0
    n_text_files = 0
    text_bytes = 0

    def _account(sfx: str, byte_size: int) -> None:
        nonlocal n_images, n_text_files, text_bytes
        if sfx in _IMAGE_EXTS:
            n_images += 1
            text_bytes += _PRE_SCAN_BINARY_TEXT_BYTES
        elif sfx in {".pdf", ".docx", ".doc"}:
            n_text_files += 1
            text_bytes += _PRE_SCAN_BINARY_TEXT_BYTES
        elif sfx in _TEXT_EXTS:
            n_text_files += 1
            text_bytes += byte_size

    for p in paths:
        path = Path(p)
        if not path.exists():
            continue
        files = [path] if path.is_file() else (
            [c for c in path.rglob("*") if c.is_file()])
        for child in files:
            sfx = child.suffix.lower()
            if sfx == ".zip":
                try:
                    with zipfile.ZipFile(child) as zf:
                        for info in zf.infolist():
                            inner_sfx = Path(info.filename).suffix.lower()
                            _account(inner_sfx, info.file_size)
                except Exception:
                    pass
            else:
                try:
                    size = child.stat().st_size
                except OSError:
                    continue
                _account(sfx, size)

    if n_topics is None:
        from engine.progress import _default_n_topics
        n_topics = _default_n_topics()
    return {
        "n_images": n_images,
        "n_text_files": n_text_files,
        "text_bytes": text_bytes,
        "n_splits": None,           # splitter hasn't run yet
        "chunk_cap": chunk_cap,
        "n_topics": n_topics,
    }


def pipeline_stats_from_docs(
    docs,
    n_splits: int | None = None,
    chunk_cap: int = 30_000,
    n_topics: int | None = None,
) -> dict:
    """Post-ingest pipeline stats. Same `PipelineStats` shape as
    `pipeline_stats_from_paths`, just sourced from parsed Documents
    + splitter output. `n_splits` is `sum(parents.values())` — the
    splitter's chunk count, which equals the real extract est_calls."""
    n_images = sum(1 for d in docs if d.source_type == SourceType.IMAGE)
    n_text_files = sum(
        1 for d in docs if d.source_type != SourceType.IMAGE)
    # Sum of REAL content lengths across all docs (image transcripts
    # + text-file content + extracted PDF/DOCX text).
    text_bytes = sum(len(d.content) for d in docs)
    if n_topics is None:
        from engine.progress import _default_n_topics
        n_topics = _default_n_topics()
    return {
        "n_images": n_images,
        "n_text_files": n_text_files,
        "text_bytes": text_bytes,
        "n_splits": n_splits,
        "chunk_cap": chunk_cap,
        "n_topics": n_topics,
    }



_TEXT_EXTS = {".txt", ".md", ".markdown", ".json", ".jsonl"}

_BINARY_SUPPORTED = {".pdf", ".docx", ".doc", ".zip"}

SUPPORTED_EXTS = _IMAGE_EXTS | _TEXT_EXTS | _BINARY_SUPPORTED

# Explicitly excluded — not "unknown", deliberately skipped
_EXCLUDED_EXTS = {".html", ".htm"}


@dataclass
class Document:
    id: str
    source_path: str
    source_type: SourceType
    content: str
    title: str = ""
    date: str = ""
    # Stable short identifier for the source file (relative path from the
    # picked input dir, or basename for single-file ingests). Unchanged by
    # segmentation — all sub-docs from the same file share the same file_id.
    # Defaults to `id` when not set explicitly (top-level ingested docs).
    file_id: str = ""
    # Offset of this document's content within the original ingested file.
    # 0 for top-level ingested docs; >0 for sub-documents created by the
    # segmenter when splitting a bundle.
    origin_char: int = 0
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        if not self.file_id:
            self.file_id = self.id


# ── Format detection ──────────────────────────────────────────────────────────

_WHATSAPP_LINE = re.compile(
    r"^\[?\d{1,2}/\d{1,2}/\d{2,4},?\s+\d{1,2}:\d{2}(:\d{2})?\s*(AM|PM)?\]?\s+.+?:"
)


# ── Text parsers ──────────────────────────────────────────────────────────────
#
# All parsers take a `file_id`: a short, stable identifier for the file that
# becomes Document.id. For files ingested directly it's the basename; for
# files inside a picked directory it's the path relative to that directory;
# for files inside a zip it's `outer.zip::inner/path`. This keeps downstream
# IDs readable (doc IDs, section IDs, chunk IDs, evidence refs) without
# embedding the user's absolute filesystem path.


def _parse_whatsapp(path: Path, file_id: str, content: str) -> list[Document]:
    content = re.sub(r"<Media omitted>\n?", "", content)
    content = re.sub(r"\u200e", "", content)  # LTR mark
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.WHATSAPP,
        content=content.strip(),
        title=path.stem,
    )]


def _parse_md_file(path: Path, file_id: str, content: str) -> list[Document]:
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.MD_FILE,
        content=content.strip(),
        title=path.stem,
    )]


def _parse_txt(path: Path, file_id: str, content: str) -> list[Document]:
    if _WHATSAPP_LINE.search(content[:2000]):
        return _parse_whatsapp(path, file_id, content)
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.TXT,
        content=content.strip(),
        title=path.stem,
    )]


# ── Day One JSON ──────────────────────────────────────────────────────────────

def _dayone_default_location(entries: list[dict]) -> str | None:
    from collections import Counter
    from datetime import datetime, timezone, timedelta

    cutoff = datetime.now(timezone.utc) - timedelta(days=90)
    counts: Counter = Counter()
    for e in entries:
        raw_date = e.get("creationDate", "")
        try:
            created = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
        except ValueError:
            continue
        if created >= cutoff:
            loc = e.get("location") or {}
            city = loc.get("localityName", "").strip()
            if city:
                counts[city] += 1
    return counts.most_common(1)[0][0] if counts else None


def _dayone_clean_text(text: str) -> str:
    text = re.sub(r"\\([^\\])", r"\1", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _dayone_local_date(creation_date: str, tz_name: str | None) -> str:
    """Return the local calendar date string (YYYY-MM-DD) for a Day One entry.

    Falls back to the raw UTC date slice when tz_name is missing or unrecognised.
    """
    from datetime import datetime, timezone
    try:
        dt_utc = datetime.fromisoformat(creation_date.replace("Z", "+00:00"))
    except (ValueError, AttributeError):
        return creation_date[:10]
    if tz_name:
        try:
            from zoneinfo import ZoneInfo, ZoneInfoNotFoundError
            return dt_utc.astimezone(ZoneInfo(tz_name)).date().isoformat()
        except (ZoneInfoNotFoundError, KeyError):
            pass
    return dt_utc.astimezone(timezone.utc).date().isoformat()


def _split_dayone_file_id(file_id: str, year_key: str) -> str:
    if "." in file_id.rsplit("::", 1)[-1]:
        stem, ext = file_id.rsplit(".", 1)
        return f"{stem}-{year_key}.{ext}"
    return f"{file_id}-{year_key}"


def _parse_dayone_json(path: Path, file_id: str, content: str) -> list[Document]:
    try:
        import json5  # type: ignore
        data = json5.loads(content)
    except ImportError:
        cleaned = re.sub(r",\s*([}\]])", r"\1", content)
        data = json.loads(cleaned)

    # Day One always serialises as a dict-with-entries; any other shape
    # (list / scalar) lands here only when the chat_exports sniff
    # declined and the JSON is genuinely something else (e.g. a
    # Claude.ai users.json roster). Silently skip rather than blow up.
    if not isinstance(data, dict):
        return []

    entries = data.get("entries", [])
    default_location = _dayone_default_location(entries)

    # Bucket text-bearing entries by year (from creationDate). Entries with
    # no parseable creationDate go into a synthetic "undated" bucket — the
    # alternative is silent data loss, which is worse than a small extra
    # source file. Empty buckets never form because entries with no text
    # are skipped before bucketing.
    from collections import OrderedDict
    buckets: "OrderedDict[str, list[tuple[dict, str]]]" = OrderedDict()
    for entry in entries:
        text = _dayone_clean_text(entry.get("text", ""))
        if not text:
            continue
        raw_date = entry.get("creationDate", "")
        local_date = _dayone_local_date(raw_date, entry.get("timeZone"))
        try:
            year_key = str(int(local_date[:4]))
        except (ValueError, TypeError):
            year_key = "undated"
        buckets.setdefault(year_key, []).append((entry, text))

    if not buckets:
        return []

    # The original filename remains the "hot" bucket new entries land in
    # (Day One re-syncs append to the original on disk), so it keeps the
    # most recent year of journaling + any undated entries. Older years
    # are pulled out into archive subfiles for navigability. Most-recent
    # year = max numeric year present; undated always rides with the
    # original regardless of whether a year bucket exists.
    year_only_keys = [k for k in buckets if k != "undated"]
    home_year_key = max(year_only_keys) if year_only_keys else None

    # Group entries by their final destination file_id. Multiple source
    # buckets (the home year + undated) can land in the same destination,
    # so we collect first and enumerate per-destination so entry indices
    # stay unique within each emitted file.
    final_groups: "OrderedDict[str, list[tuple[dict, str, bool]]]" = OrderedDict()
    for year_key in sorted(buckets.keys(), key=lambda k: (k == "undated", k)):
        is_home = (year_key == home_year_key) or (year_key == "undated")
        dest_file_id = file_id if is_home else _split_dayone_file_id(file_id, year_key)
        for entry, text in buckets[year_key]:
            final_groups.setdefault(dest_file_id, []).append((entry, text, is_home))

    docs = []
    for dest_file_id, items in final_groups.items():
        for entry_idx, (entry, text, is_home) in enumerate(items):
            date = _dayone_local_date(entry.get("creationDate", ""), entry.get("timeZone"))
            uuid = entry.get("uuid", "")
            tags = entry.get("tags", [])

            loc = entry.get("location") or {}
            city = loc.get("localityName", "").strip()
            country = loc.get("country", "").strip()
            location = f"{city}, {country}".strip(", ") if city else None
            if location and city == default_location:
                location = None

            # Full Day One UUID is preserved in metadata for reverse lookups;
            # the id uses a short entry index so logs and intermediate paths
            # stay readable.
            metadata: dict = {"uuid": uuid}
            if tags:
                metadata["tags"] = tags
            if location:
                metadata["location"] = location
            if not is_home:
                # Traceability back to the pre-split source so a fact's
                # evidence anchor on a split-out historical year can still
                # surface "originally from <X>".
                metadata["origin_file_id"] = file_id

            docs.append(Document(
                id=f"{dest_file_id}::entry_{entry_idx:03d}",
                source_path=str(path),
                source_type=SourceType.DAYONE_JSON,
                content=text,
                title=f"Journal {date}",
                date=date,
                file_id=dest_file_id,
                metadata=metadata,
            ))
    return docs


# ── Binary parsers ────────────────────────────────────────────────────────────

def _parse_pdf(path: Path, file_id: str) -> list[Document]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        _log(f"  Warning: failed to read PDF {path}: {e}")
        return []
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    content = "\n\n".join(pages).strip()
    if not content:
        _log(f"  PDF has no extractable text (likely scanned), skipping: {path}")
        return []
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.PDF,
        content=content,
        title=path.stem,
    )]


def _parse_docx(path: Path, file_id: str) -> list[Document]:
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        _log(f"  Warning: failed to read DOCX {path}: {e}")
        return []
    content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not content.strip():
        return []
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.DOCX,
        content=content.strip(),
        title=path.stem,
    )]


def _parse_image(path: Path, file_id: str, vision_mode: Mode = Mode.TEE) -> list[Document]:
    """Emit a placeholder image Document; defer the actual vision
    transcription to the runner's dedicated vision stage
    (`vision.describe_images_all`).

    Pre-refactor, this parser invoked `describe_image` synchronously
    inside the ingestor's per-file ThreadPoolExecutor — which made
    vision calls fan out concurrently across N worker threads. The
    scheduler can't pace concurrent same-stage dispatches via
    `run_stage` (exclusive `_stage_active` lock), so vision was
    forced into a bypass that dropped the gate entirely. Lifting
    dispatch into a runner-level stage with one
    `scheduler.run_stage("vision", producer)` call across all
    image docs restores pacing without re-introducing the
    concurrency bug; matches the shape of every other stage.

    The Document carries `metadata["_pending_vision"] = True` and
    `metadata["_vision_mode"] = vision_mode.value` so the runner's
    vision stage can identify which docs need transcription and
    under which mode (the vision mode must match the pipeline's
    trust boundary; the runner threads the same value the ingestor
    received).
    """
    return [Document(
        id=file_id,
        source_path=str(path),
        source_type=SourceType.IMAGE,
        content="",
        title=path.stem,
        metadata={
            "_pending_vision": True,
            "_vision_mode": vision_mode.value,
        },
    )]


# ── Zip handler — recursive ───────────────────────────────────────────────────

def _parse_zip(path: Path, file_id: str, vision_mode: Mode = Mode.TEE) -> list[Document]:
    """
    Extract a zip and recursively ingest every file inside, in parallel.
    Nested zips are handled via recursion through _ingest_file.

    Inner file IDs are constructed as `{zip_file_id}::{relative_inner_path}`.
    """
    docs = []
    tmpdir = tempfile.mkdtemp(prefix="bvvault_zip_")
    try:
        tmppath = Path(tmpdir)
        try:
            with zipfile.ZipFile(path) as zf:
                zf.extractall(tmppath)
        except Exception as e:
            _log(f"  Failed to extract zip {path}: {e}")
            return []

        children = [
            c for c in sorted(tmppath.rglob("*"))
            if c.is_file()
            and not c.name.startswith("._")
            and "__MACOSX" not in c.parts
        ]

        if not children:
            return []

        def _process(child: Path) -> list[Document]:
            rel = child.relative_to(tmppath)
            inner_file_id = f"{file_id}::{rel}"
            inner_docs = _ingest_file(child, vision_mode=vision_mode, file_id=inner_file_id)
            for d in inner_docs:
                d.source_path = f"{path}::{rel}"
                d.metadata.setdefault("origin_zip", str(path))
            return inner_docs

        n_workers = min(len(children), max_workers(vision_mode))
        with ThreadPoolExecutor(max_workers=n_workers) as ex:
            futures = [ex.submit(_process, c) for c in children]
            for f in as_completed(futures):
                docs.extend(f.result())
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    return docs


# ── File dispatcher ───────────────────────────────────────────────────────────

def _ingest_file(
    path: Path,
    vision_mode: Mode = Mode.TEE,
    file_id: str | None = None,
) -> list[Document]:
    """
    Parse one file into Documents. `file_id` is a short stable identifier
    (e.g. basename or relative path from the picked input dir) that becomes
    Document.id. Defaults to the file basename.
    """
    suffix = path.suffix.lower()
    if file_id is None:
        file_id = path.name

    # Size gate — drop files larger than MAX_FILE_SIZE
    try:
        size = path.stat().st_size
    except OSError:
        _log(f"  Cannot stat {path}, skipping")
        return []
    if size > MAX_FILE_SIZE:
        _log(f"  Skipping {path.name}: {size / 1024 / 1024:.1f} MB > 20 MB limit")
        return []

    # Explicit exclusion
    if suffix in _EXCLUDED_EXTS:
        _log(f"  Skipping excluded format: {path.name}")
        return []

    # Whitelist — everything else drops
    if suffix not in SUPPORTED_EXTS:
        _log(f"  Skipping unsupported format: {path.name}")
        return []

    # Dispatch
    if suffix == ".zip":
        return _parse_zip(path, file_id, vision_mode=vision_mode)

    if suffix == ".pdf":
        return _parse_pdf(path, file_id)

    if suffix in (".docx", ".doc"):
        return _parse_docx(path, file_id)

    if suffix in _IMAGE_EXTS:
        return _parse_image(path, file_id, vision_mode=vision_mode)

    # Text-based formats
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        _log(f"  Could not read {path}: {e}")
        return []

    if suffix in (".md", ".markdown"):
        return _parse_md_file(path, file_id, content)
    if suffix == ".txt":
        return _parse_txt(path, file_id, content)
    if suffix == ".json":
        # Shape-sniff: chat-export parsers get first crack; Day One is
        # the fallback so existing journal exports keep working with
        # no change. An unrecognized JSON shape drops silently — same
        # behaviour the Day One parser produced before for non-journal
        # JSON (returned [] on missing `entries`).
        from engine.chat_exports import parse_chat_json
        chat_docs = parse_chat_json(path, file_id, content)
        if chat_docs is not None:
            return chat_docs
        return _parse_dayone_json(path, file_id, content)
    if suffix == ".jsonl":
        from engine.chat_exports import parse_chat_jsonl
        return parse_chat_jsonl(path, file_id, content)

    # Should not reach here given SUPPORTED_EXTS whitelist above
    return []


def _ingest_dir(path: Path, vision_mode: Mode = Mode.TEE) -> list[Document]:
    """Ingest every file in a directory tree, in parallel.

    Each file's `file_id` is its path relative to the selected input dir,
    so a file at `data/daily_insights/insights_2025-12-08.txt` (when `data/`
    is picked) gets file_id `daily_insights/insights_2025-12-08.txt` — no
    absolute path leakage, but subdirectory context preserved.
    """
    children = [c for c in sorted(path.rglob("*")) if c.is_file()]
    if not children:
        return []

    def _one(child: Path) -> list[Document]:
        try:
            rel = child.relative_to(path)
            file_id = str(rel)
        except ValueError:
            file_id = child.name
        return _ingest_file(child, vision_mode=vision_mode, file_id=file_id)

    docs = []
    n_workers = min(len(children), max_workers(vision_mode))
    with ThreadPoolExecutor(max_workers=n_workers) as ex:
        futures = [ex.submit(_one, c) for c in children]
        for f in as_completed(futures):
            docs.extend(f.result())
    return docs


# ── Public API ────────────────────────────────────────────────────────────────

def ingest(input_path: str | Path, vision_mode: Mode = Mode.TEE) -> list[Document]:
    """
    Ingest a file or directory. Returns a list of Document objects.

    vision_mode controls which LLM provider transcribes images. Defaults
    to Mode.TEE (the attested production route).
    """
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"Input not found: {path}")

    if path.is_dir():
        docs = _ingest_dir(path, vision_mode=vision_mode)
    else:
        docs = _ingest_file(path, vision_mode=vision_mode)

    _log(f"Ingested {len(docs)} document(s) from {path}")
    return docs
